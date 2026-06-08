# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Inches, Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(HERE, "img")
OUT = os.path.normpath(os.path.join(HERE, "..", "BadmintonHub-UseCase-Flowchart.docx"))

FONT = "Arial"
LABELS = ["Mo ta", "Dau vao", "Xu ly", "Ket qua"]  # replaced below w/ unicode
LABELS = ["Mô tả", "Đầu vào", "Xử lý", "Kết quả"]

fig_no = 0

def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def style_run(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(a), FONT)
    if color is not None:
        run.font.color.rgb = color

def add_para(doc, text, size=11, bold=False, align=None, italic=False,
             space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    run.italic = italic
    style_run(run, size=size, bold=bold, color=color)
    return p

def add_heading(doc, text, level):
    sizes = {1: 16, 2: 14, 3: 12.5, 4: 11.5}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=sizes.get(level, 11), bold=True,
              color=RGBColor(0x1F, 0x2A, 0x44))
    return p

def fit(path, max_w=6.0, max_h=7.6):
    with Image.open(path) as im:
        w, h = im.size
    ar = w / float(h)
    w_in = max_w
    h_in = w_in / ar
    if h_in > max_h:
        h_in = max_h
        w_in = h_in * ar
    return Inches(w_in), Inches(h_in)

def add_figure(doc, filename, caption, max_w=6.0, max_h=7.6):
    global fig_no
    fig_no += 1
    path = os.path.join(IMG, filename)
    w, h = fit(path, max_w, max_h)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=w, height=h)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run("Hình %d. %s" % (fig_no, caption))
    style_run(run, size=10, bold=False)
    run.italic = True

def add_spec_table(doc, desc, inp, proc, result):
    rows = [desc, inp, proc, result]
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, content in enumerate(rows):
        lbl = table.cell(i, 0)
        val = table.cell(i, 1)
        set_cell_width(lbl, 1.35)
        set_cell_width(val, 5.15)
        set_cell_bg(lbl, "E8ECF4")
        lp = lbl.paragraphs[0]
        lp.paragraph_format.space_after = Pt(2)
        lp.paragraph_format.space_before = Pt(2)
        lr = lp.add_run(LABELS[i])
        style_run(lr, size=11, bold=True)
        if isinstance(content, list):
            for j, line in enumerate(content):
                vp = val.paragraphs[0] if j == 0 else val.add_paragraph()
                vp.paragraph_format.space_after = Pt(1)
                vp.paragraph_format.space_before = Pt(1)
                vr = vp.add_run(line)
                style_run(vr, size=11)
        else:
            vp = val.paragraphs[0]
            vp.paragraph_format.space_after = Pt(2)
            vp.paragraph_format.space_before = Pt(2)
            vr = vp.add_run(content)
            style_run(vr, size=11)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_function(doc, title, image, desc, inp, proc, result):
    add_heading(doc, title, 4)
    add_figure(doc, image, title)
    add_spec_table(doc, desc, inp, proc, result)

# ---------------------------------------------------------------- document
doc = Document()
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Mm(20)
sec.bottom_margin = Mm(20)
sec.left_margin = Mm(22)
sec.right_margin = Mm(22)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(24)
tr = t.add_run("HỆ THỐNG BADMINTONHUB")
style_run(tr, size=22, bold=True, color=RGBColor(0x1F, 0x2A, 0x44))
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
st.paragraph_format.space_after = Pt(18)
sr = st.add_run("Đặc tả Use Case và Sơ đồ luồng xử lý (Flowchart)")
style_run(sr, size=14, bold=False, color=RGBColor(0x44, 0x4A, 0x55))
sr.italic = True

# ============================================ 2. Customer Requirements
add_heading(doc, "2. Đặc tả yêu cầu người dùng", 1)

add_heading(doc, "2.1. Người dùng của hệ thống", 2)
add_para(doc, "Hệ thống BadmintonHub phục vụ ba nhóm người dùng chính, mỗi nhóm có nhu cầu và mức độ tương tác khác nhau:")
add_para(doc, "Khách hàng (Customer): nhóm người dùng cuối, bao gồm khách vãng lai (chưa đăng nhập – tra cứu sân, sản phẩm, đăng ký tài khoản) và khách đã đăng nhập. Họ đặt sân theo lượt, đăng ký lịch cố định theo gói, mua sản phẩm, thanh toán trực tuyến, đánh giá sân và xem lịch sử giao dịch.")
add_para(doc, "Nhân viên (Employee): lực lượng vận hành trực tiếp tại cơ sở. Họ check-in khách tại sân, đối soát và xác nhận các thanh toán thủ công, bán hàng tại quầy (POS) và nhập kho.")
add_para(doc, "Quản trị viên (Admin): nhóm có quyền cao nhất. Họ duyệt đơn bán, quản lý đơn mua hàng từ nhà cung cấp, điều chuyển kho, quản lý sân/sản phẩm/giá, người dùng & phân quyền và theo dõi báo cáo.")
add_para(doc, "Mối quan hệ: khách hàng tương tác trực tiếp qua web/app; nhân viên là cầu nối xử lý các tình huống tại chỗ; quản trị viên giám sát và điều hành toàn bộ hệ thống.")

add_heading(doc, "2.2. Sơ đồ Use Case", 2)
add_para(doc, "Các sơ đồ use case dưới đây mô tả tập hợp chức năng (use case) mà mỗi nhóm người dùng có thể thực hiện trong phạm vi hệ thống.")

add_heading(doc, "2.2.1. Use Case – Khách hàng", 3)
add_figure(doc, "uc-1.png", "Use Case – Khách hàng", max_w=6.2, max_h=7.2)

add_heading(doc, "2.2.2. Use Case – Nhân viên", 3)
add_figure(doc, "uc-2.png", "Use Case – Nhân viên", max_w=6.2, max_h=7.2)

add_heading(doc, "2.2.3. Use Case – Quản trị viên", 3)
add_figure(doc, "uc-3.png", "Use Case – Quản trị viên", max_w=6.2, max_h=7.2)

# ============================================ 3. System Functions Design
add_heading(doc, "3. Thiết kế chức năng hệ thống", 1)
add_para(doc, "Mỗi chức năng được mô tả bằng một sơ đồ luồng xử lý (flowchart) kèm bảng đặc tả gồm: Mô tả, Đầu vào, Xử lý và Kết quả.")

# ---- 3.1 General
add_heading(doc, "3.1. Chức năng chung", 2)
add_function(
    doc, "Đăng nhập hệ thống", "flow-1.png",
    "Cho phép người dùng (Khách hàng, Nhân viên, Quản trị) xác thực danh tính để truy cập hệ thống theo đúng vai trò.",
    "Tên đăng nhập (email/số điện thoại) và mật khẩu.",
    ["1. Người dùng nhập thông tin tại form đăng nhập.",
     "2. Hệ thống kiểm tra thông tin hợp lệ.",
     "3. Nếu sai: báo lỗi và quay lại form đăng nhập.",
     "4. Nếu đúng: phát hành JWT token và nạp thông tin người dùng theo vai trò."],
    "Thành công: cấp JWT và vào hệ thống theo phân quyền. Thất bại: thông báo sai thông tin đăng nhập.")

# ---- 3.2 Customer
add_heading(doc, "3.2. Chức năng Khách hàng", 2)
add_function(
    doc, "Đặt sân theo lượt", "flow-2.png",
    "Cho phép khách đặt một khung giờ của sân trong ngày. Đặt sân bắt buộc thanh toán trực tuyến (không thu tiền mặt).",
    "Sân, ngày và khung giờ mong muốn.",
    ["1. Khách chọn sân, ngày và khung giờ.",
     "2. Hệ thống kiểm tra slot khả dụng.",
     "3. Nếu đã có người đặt: yêu cầu chọn lại.",
     "4. Nếu còn trống: tạo booking trạng thái chờ và khóa slot.",
     "5. Sinh hóa đơn chưa thanh toán.",
     "6. Chuyển sang thanh toán trực tuyến bắt buộc."],
    "Booking được giữ chỗ kèm hóa đơn; khách thanh toán trực tuyến để xác nhận.")
add_function(
    doc, "Đặt lịch cố định theo gói", "flow-3.png",
    "Cho phép khách đăng ký chơi định kỳ trong một khoảng ngày dưới dạng một gói gồm nhiều buổi.",
    "Sân, chu kỳ lặp và khoảng ngày.",
    ["1. Hệ thống tạo danh sách buổi dự kiến (preview).",
     "2. Kiểm tra tất cả buổi có trống không.",
     "3. Nếu có buổi trùng: hiển thị để đổi giờ/đổi sân và tạo lại preview.",
     "4. Khi tất cả trống: tính giá gói và chốt thông tin khách.",
     "5. Tạo gói cố định cùng các buổi và booking.",
     "6. Thanh toán trực tuyến cho cả gói."],
    "Gói cố định cùng toàn bộ buổi được tạo và thanh toán trực tuyến.")
add_function(
    doc, "Đặt mua sản phẩm", "flow-4.png",
    "Cho phép khách mua sản phẩm (vợt, cầu, phụ kiện) giao tận nơi hoặc nhận tại cửa hàng.",
    "Sản phẩm trong giỏ, thông tin nhận hàng và hình thức nhận.",
    ["1. Khách chọn sản phẩm vào giỏ hàng.",
     "2. Nhập thông tin nhận hàng và chọn hình thức nhận.",
     "3. Hệ thống chọn kho phục vụ và kiểm tra tồn.",
     "4. Nếu thiếu tồn: báo sản phẩm hết hàng.",
     "5. Nếu đủ: tạo đơn chờ xử lý và sinh hóa đơn.",
     "6. Thanh toán trực tuyến hoặc tiền mặt khi nhận (COD/nhận tại cửa hàng)."],
    "Đơn hàng được tạo kèm hóa đơn; thanh toán online hoặc tiền mặt khi nhận.")
add_function(
    doc, "Thanh toán", "flow-5.png",
    "Xử lý thanh toán hóa đơn. Phương thức tự xác thực được thì hệ thống tự xác nhận; ngược lại nhân viên đối soát thủ công.",
    "Hóa đơn cần thanh toán và phương thức thanh toán.",
    ["1. Khách chọn phương thức thanh toán.",
     "2. Nếu cổng tự xác thực được (VNPay/MoMo/SePay): tạo link/QR, khách thanh toán, cổng gửi IPN để hệ thống tự xác nhận.",
     "3. Nếu không (chuyển khoản thủ công/tiền mặt): nhân viên đối soát và xác nhận.",
     "4. Cập nhật hóa đơn đã thanh toán."],
    "Hóa đơn chuyển trạng thái đã thanh toán; giao dịch tự xác nhận hoặc do nhân viên xác nhận.")
add_function(
    doc, "Đánh giá sân", "flow-6.png",
    "Cho phép khách đã chơi đánh giá sân bằng số sao và nội dung; mỗi sân chỉ đánh giá một lần.",
    "Sân đã từng chơi, số sao và nội dung đánh giá.",
    ["1. Khách chọn sân đã chơi.",
     "2. Nhập số sao và nội dung.",
     "3. Kiểm tra đã đánh giá sân này chưa.",
     "4. Nếu rồi: báo không cho gửi trùng.",
     "5. Nếu chưa: lưu đánh giá và cập nhật điểm trung bình của sân."],
    "Đánh giá được lưu và điểm trung bình của sân được cập nhật.")

# ---- 3.3 Employee
add_heading(doc, "3.3. Chức năng Nhân viên", 2)
add_function(
    doc, "Check-in khách tại sân", "flow-7.png",
    "Nhân viên xác nhận khách đến và cho vào sân dựa trên mã booking.",
    "Mã booking (quét QR hoặc nhập tay).",
    ["1. Nhân viên quét/nhập mã booking.",
     "2. Hệ thống tra cứu thông tin booking.",
     "3. Kiểm tra hợp lệ và đã thanh toán.",
     "4. Nếu chưa: báo lỗi hoặc yêu cầu thanh toán trước.",
     "5. Nếu hợp lệ: cập nhật trạng thái sang đang chơi và cho khách vào sân."],
    "Booking chuyển sang đang chơi; khách được vào sân.")
add_function(
    doc, "Xác nhận thanh toán thủ công", "flow-8.png",
    "Nhân viên đối soát và xác nhận các thanh toán không tự xác thực (chuyển khoản/tiền mặt).",
    "Booking/hóa đơn đang chờ thanh toán.",
    ["1. Nhân viên mở booking đang chờ thanh toán.",
     "2. Đối soát chuyển khoản hoặc tiền mặt.",
     "3. Kiểm tra khớp số tiền.",
     "4. Nếu chưa khớp: giữ trạng thái chờ và ghi chú.",
     "5. Nếu khớp: cập nhật hóa đơn đã thanh toán và chuyển booking sang đã xác nhận."],
    "Hóa đơn đã thanh toán, booking được xác nhận; hoặc giữ chờ kèm ghi chú nếu lệch.")
add_function(
    doc, "Bán hàng tại quầy (POS)", "flow-9.png",
    "Nhân viên lập đơn bán tại quầy cho khách (kể cả khách vãng lai) và gửi quản lý duyệt.",
    "Khách hàng (hoặc khách vãng lai) và danh sách sản phẩm.",
    ["1. Nhân viên chọn khách hoặc tạo khách vãng lai.",
     "2. Thêm sản phẩm vào đơn bán.",
     "3. Tạo đơn bán trạng thái chờ duyệt.",
     "4. Gửi quản lý phê duyệt."],
    "Đơn bán ở trạng thái chờ duyệt, chờ quản lý phê duyệt.")
add_function(
    doc, "Nhập kho", "flow-10.png",
    "Nhân viên lập phiếu nhập để tăng tồn kho cho sản phẩm tại một kho.",
    "Kho, sản phẩm và số lượng nhập.",
    ["1. Nhân viên mở phiếu nhập kho.",
     "2. Chọn kho, sản phẩm và nhập số lượng.",
     "3. Kiểm tra thông tin hợp lệ.",
     "4. Nếu sai: báo lỗi và nhập lại.",
     "5. Nếu hợp lệ: cộng tồn on_hand và available, ghi nhận giao dịch nhập kho."],
    "Tồn kho được cộng và giao dịch nhập kho được ghi nhận.")

# ---- 3.4 Admin
add_heading(doc, "3.4. Chức năng Quản trị", 2)
add_function(
    doc, "Duyệt đơn bán", "flow-11.png",
    "Quản lý duyệt đơn bán POS, xác nhận thanh toán và xuất kho.",
    "Đơn bán đang chờ duyệt.",
    ["1. Quản lý mở đơn bán chờ duyệt.",
     "2. Kiểm tra thông tin đơn và tồn kho.",
     "3. Nếu không duyệt: từ chối đơn và lưu lý do.",
     "4. Nếu duyệt: cập nhật trạng thái đã duyệt, xác nhận thanh toán, xuất kho và trừ tồn."],
    "Đơn bán được duyệt, thanh toán xác nhận và trừ tồn; hoặc bị từ chối kèm lý do.")
add_function(
    doc, "Đơn mua hàng nhà cung cấp", "flow-12.png",
    "Quản trị tạo đơn mua hàng từ nhà cung cấp và nhập kho khi nhận hàng.",
    "Nhà cung cấp và danh sách sản phẩm cần mua.",
    ["1. Tạo đơn mua hàng trạng thái nháp.",
     "2. Gửi đơn cho nhà cung cấp.",
     "3. Nếu nhà cung cấp không xác nhận: hủy đơn mua hàng.",
     "4. Nếu xác nhận: nhận hàng và lập phiếu nhập kho, cập nhật tồn, đơn chuyển đã nhận."],
    "Hàng được nhập kho và đơn mua chuyển trạng thái đã nhận; hoặc đơn bị hủy.")
add_function(
    doc, "Chuyển kho", "flow-13.png",
    "Quản trị điều chuyển tồn giữa hai kho và cân bằng tồn khi hoàn tất.",
    "Kho nguồn, kho đích, sản phẩm và số lượng chuyển.",
    ["1. Tạo yêu cầu chuyển kho từ kho A sang kho B.",
     "2. Kiểm tra yêu cầu được duyệt không.",
     "3. Nếu không: từ chối yêu cầu.",
     "4. Nếu duyệt: xuất kho nguồn (đang chuyển), nhập kho đích khi hàng tới, hoàn tất và cân bằng tồn hai kho."],
    "Tồn được chuyển và cân bằng giữa hai kho; hoặc yêu cầu bị từ chối.")
add_function(
    doc, "Dashboard và báo cáo", "flow-14.png",
    "Quản trị xem thống kê doanh thu, booking và đơn hàng theo khoảng thời gian.",
    "Khoảng thời gian thống kê.",
    ["1. Quản trị chọn khoảng thời gian.",
     "2. Hệ thống tổng hợp doanh thu, booking và đơn hàng.",
     "3. Nếu không có dữ liệu: hiển thị trạng thái trống.",
     "4. Nếu có: hiển thị biểu đồ và các chỉ số chính."],
    "Bảng điều khiển hiển thị biểu đồ và chỉ số chính của kỳ đã chọn.")

doc.save(OUT)
print("SAVED:", OUT)
print("Figures:", fig_no)
